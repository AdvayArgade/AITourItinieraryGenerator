from datetime import timedelta
import json
import math
import requests
import streamlit as st
from geopy.geocoders import Nominatim
from openai import OpenAI
import logging
from dotenv import load_dotenv
import os
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO


load_dotenv('keys.env')
API_KEY = os.getenv('OPENAI_API_KEY')
RAPID_API_KEY = os.getenv("X-RapidAPI-Key")
RAPID_API_HOST = os.getenv("X-RapidAPI-Host")


client = OpenAI(api_key=API_KEY)

function_descriptions = [
    {
        "name": "get_flight_hotel_info",
        "description": "Find the flights between cities and hotels within cities for residence.",
        "parameters": {
            "type": "object",
            "properties": {
                "loc_list": {
                  "type": "array",
                  "items": {
                      "type": "string"
                  },
                  "description": "The ordered list of names of cities in the tour. e.g. ['Mumbai', 'Paris']"
              },

                "date_list": {
                  "type": "array",
                  "items": {
                      "type": "string"
                  },
                  "description": "The ordered list of dates for arrival in the cities in YYYY-MM-DD format."
              },

            },
            "required": ["loc_list", "date_list"],
        },
    }
]

# Set the logging level (DEBUG for most details)
logging.basicConfig(level=logging.DEBUG)

# Create a logger for your specific API usage
logger = logging.getLogger('booking_com_api')

def make_booking_com_api_call(url, headers, data):
  logger.debug("API Request:")
  logger.debug(f"URL: {url}")
  logger.debug(f"Headers: {headers}")
  logger.debug(f"Body: {data}")

  # Make your API call using requests library
  response = requests.get(url, headers=headers, json=data)

  logger.debug("API Response:")
  logger.debug(f"Status Code: {response.status_code}")
  logger.debug(f"Headers: {response.headers}")
  logger.debug(f"Response Body: {response.json()}")

  return response

def get_hotel_data(city, checkin_date, checkout_date, num_adults, num_children):
    city_dict = {}

    # Geocode city to get latitude and longitude
    geolocator = Nominatim(user_agent="MyApp")
    location = geolocator.geocode(city)

    if location:
        lat = location.latitude
        long = location.longitude

        # Define URL and query parameters for Booking.com API
        url = "https://booking-com.p.rapidapi.com/v1/hotels/search-by-coordinates"
        num_rooms = math.ceil((num_adults + num_children) / 3)
        querystring = {
            "locale": "en-gb",
            "room_number": num_rooms,
            "checkin_date": checkin_date,
            "checkout_date": checkout_date,
            "filter_by_currency": "INR",
            "longitude": long,
            "latitude": lat,
            "adults_number": num_adults,
            "order_by": "popularity",
            "units": "metric",
            "page_number": "0",
        }
        if num_children>0:
            querystring["children_number"] = num_children
        headers = {
            "X-RapidAPI-Key": RAPID_API_KEY,
            "X-RapidAPI-Host": RAPID_API_HOST
        }

        # Send request to Booking.com API
        response = requests.get(url, headers=headers, params=querystring)
        if response.status_code == 200:
            data = response.json()
            if "result" in data:
                city_dict[city] = []
                st.write(f"Finding hotels for {city}")
                # st.write("Hotel search results for", city, ":")
                for index, hotel in enumerate(data["result"], start=1):
                    hotel_dict = {}
                    hotel_dict['hotel_name'] = hotel.get("hotel_name", "N/A")
                    price = hotel.get("min_total_price", "N/A")
                    if price is not None:
                        hotel_dict['price'] = str(price)
                    else:
                        hotel_dict['price'] = "N/A"
                    hotel_dict['address'] = hotel.get("address", "N/A")
                    hotel_dict['rating'] = hotel.get("review_score", "N/A")
                    # st.write(f"{index}. {hotel_dict['hotel_name']} - Address: {hotel_dict['address']}, Price for one day: {hotel_dict['price']} INR, Rating: {hotel_dict['rating']}")
                    city_dict[city].append(hotel_dict)
            else:
                st.write(f"No hotel results found for {city}.")
        else:
            st.write(f"Failed to retrieve hotel search results for {city}.")
    else:
        st.write(f"Invalid city name: {city}")

    return city_dict


def generate_itinerary(input_dict):
    # Part 1: generate the list of cities and get the hotels
    # Call the OpenAI API for creating the list of cities and dates
    input_dict['end_date'] = str(input_dict['start_date'] + timedelta(days=input_dict['num_days']))
    user_prompt = f"Generate a list of cities for a tour of {input_dict['dest']} for {input_dict['num_tourists']} " \
                  f"people with {input_dict['num_adults']} adults, purpose as {input_dict['genre']} " \
                  f"for {input_dict['num_days']} days and a budget per person of {input_dict['price_per_person']} INR starting on " \
                  f"{input_dict['start_date']}, ending on {input_dict['end_date']}. Call the function 'get_flight_hotel_info'"

    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": user_prompt}],
        # Add function calling
        functions=function_descriptions,
        function_call="auto",  # specify the function call
        max_tokens=200
    )
    output = completion.choices[0].message
    cities = json.loads(output.function_call.arguments).get("loc_list")
    dates = json.loads(output.function_call.arguments).get("date_list")
    dates.append(input_dict['end_date'])

    input_dict['cities'] = cities
    all_city_dict = {}

    city_string = ''
    for city in cities:
        city_string+=city + '  '
    st.subheader("Cities: ")
    st.write(city_string)

    for i in range(len(cities)):
        # st.write(cities[i], dates[i], dates[i+1], input_dict['num_adults'], input_dict['num_children'])
        all_city_dict.update(get_hotel_data(cities[i], dates[i], dates[i+1], input_dict['num_adults'], input_dict['num_children']))
    input_dict['hotels_by_city'] = all_city_dict


    # Part 2: Actually generate the itinerary
    user_message = f"Design a detailed itinerary for a trip from {input_dict['src']} to {input_dict['dest']} starting from {input_dict['start_date']} and for " \
                   f"{input_dict['num_days']} days. The ordered list of cities is {cities} and of dates is {dates}. The budget for this trip is {input_dict['price_per_person']} INR per person. This trip is designed " \
                   f"for {input_dict['num_tourists']} mainly with their {input_dict['type_of_travelers']} with an average age of {input_dict['average_age']}.The " \
                   f"primary interests for activities are {input_dict['genre']}.The preferred mode(s) of travel include " \
                   f"{input_dict['mode_of_travel']}.The group prefers {input_dict['food']} food. Please structure the itinerary with a detailed " \
                   f"plan for each day, including activities, locations, weather according to the season they are " \
                   f"travelling and estimated travel distances and times. Write the travel time and distance in the day's subheading. " \
                   f"Ensure to consider the preferences and " \
                   f"interests of the group for each day's schedule. Important considerations: Factor in travel time " \
                   f"between destinations. Suggest local transportation options. Include a mix of activities that cater" \
                   f" to the group's interests. Also add distance of travel for each day and approx time " \
                   f"of travel. Also you can give a name for each day in the itinerary which will be more " \
                   f"appealing. Keep the response descriptive and . Give a title to the itinerary. Do not suggest any activities " \
                   f"in the first city."

    # Generate the travel itinerary using the modified user message
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": user_message,
            }
        ],
        model="gpt-3.5-turbo",
        stream=True,
    )
    st.subheader("Itinerary")
    response = st.write_stream(chat_completion)
    # response = "Dummy response"
    return response, all_city_dict


def text_to_doc(itinerary, input_dict):
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture("logo.png", width=Inches(2.0))  # Adjust width as needed
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Split the itinerary into lines
    lines = itinerary.split('\n')

    # Add the first line as a centered and bold paragraph
    first_line = lines[0]
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(first_line)
    run.bold = True
    run.font.size = Pt(16)  # Set the font size to 16 points for "very bold"

    # Set paragraph alignment to center
    # first_line = itinerary.split('\n')[1]

    # # Add the first line as a centered header
    # header = document.add_heading(level=1)
    # header_run = header.add_run(first_line)
    # header_run.font.size = Pt(16)  # Adjust font size if needed
    # header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph1 = document.add_paragraph()
    run1 = paragraph1.add_run("Destination: ")
    run1.bold = True
    run1 = paragraph1.add_run(input_dict['dest'])

    paragraph2 = document.add_paragraph()
    run2 = paragraph2.add_run("Duration: ")
    run2.bold = True
    run2 = paragraph2.add_run(str(input_dict['num_days']))

    paragraph3 = document.add_paragraph()
    run3 = paragraph3.add_run("Interests: ")
    run3.bold = True
    run3 = paragraph3.add_run(input_dict['genre'])
    paragraph2 = document.add_paragraph()
    run2 = paragraph2.add_run("Commences on: ")
    run2.bold = True
    run2 = paragraph2.add_run(str(input_dict['start_date']))

    for line in itinerary.split('\n')[1:]:
        paragraph = document.add_paragraph()
        for char in line:
            if char == '*':
                run = paragraph.add_run()
                run.bold = True
            else:
                run = paragraph.add_run(char)


    # we will put over data here
    table_data = [
        ["Destination", " Hotel"],
    ]
    for city in input_dict['cities']:
        table_data.append([city, ''])

    table = document.add_table(rows=len(table_data), cols=2)

    # adding data to table
    for i, row_data in enumerate(table_data):
        for j, cell_data in enumerate(row_data):
            table.cell(i, j).text = cell_data

    # Apply alignment to the table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Define a custom table style (optional)
    table.style = 'Table Grid'

    return document


st.cache(suppress_st_warning=True)(get_hotel_data)
st.cache(suppress_st_warning=True)(generate_itinerary)
past_itineraries = []  # List to store past itineraries
input_dict = {}
st.set_page_config(
    page_title="AI Tour Itinerary Generator",  # Set your desired title here
    page_icon="images/favicon.ico",  # Set path to your favicon image (.ico format)
)
st.title("Tour Itinerary Generator")

col1, col2 = st.columns(2)

input_dict['dest'] = col1.text_input("Destination", key='dest')
input_dict['src'] = col1.text_input("Source City", key='src')
input_dict['genre'] = col1.text_input("Genre", key='genre')
input_dict['type_of_travelers'] = col1.text_input("Type of Travelers", key='type', placeholder='ex. family, friends')
input_dict['mode_of_travel'] = col1.text_input("Mode of Travel", key='mode', placeholder='ex. flight, bus, train')
input_dict['num_days'] = col2.number_input("Number of Days", key='num_days', min_value=0, max_value=None, value=0, step=1, format="%d")
input_dict['start_date'] = col2.date_input("Start Date", key='start_date')
# Create sub-columns within col2
col21, col22 = col2.columns(2)

input_dict['num_adults'] = int(col21.number_input("Number of Adults", key='num_adults', min_value=0, max_value=None, value=0, step=1, format="%d"))
input_dict['num_children'] = int(col22.number_input("Number of Children", key='num_children', min_value=0, max_value=None, value=0, step=1, format="%d"))
input_dict['price_per_person'] = col2.number_input("Price Per Person", key='price_per_person', min_value=0.0)
input_dict['average_age'] = col2.number_input("Average age", key='average_age', min_value=0, max_value=None, value=0, step=1, format="%d")
input_dict['food'] = 'non veg' if st.toggle('Include non-veg hotels') else 'veg'

input_dict['num_tourists'] = input_dict['num_adults'] + input_dict['num_children']

if st.button("Generate Itinerary", type="primary"):
    for key in input_dict.keys():
        if input_dict[key] is None:
            st.warning(f'Please enter {key}!')

    else:
        generated_itinerary, city_dict = generate_itinerary(input_dict)
        past_itineraries.append(generated_itinerary)  # Add to past itineraries
        isGenerated = True
        doc = text_to_doc(generated_itinerary, input_dict)
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        st.download_button(
            label="Download Word Document",
            data=doc_io,
            file_name=f"{input_dict['dest']} Itinerary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

         # Main side panel for hotels
        st.subheader("Hotels")
        for city, hotels in city_dict.items():
            city_expander = st.expander(f"{city}")
            with city_expander:
                for hotel in hotels:
                    st.write(f"- {hotel['hotel_name']}")
                    st.write(f"  Address: {hotel['address']}")
                    st.write(f"  Price per day: {hotel['price']} INR")
                    st.write(f"  Rating: {hotel['rating']}")
                    # Add more details as needed (amenities, images, etc.)
                    st.write("---")  # Separator between hotels



with st.expander("Past Itineraries"):
    if past_itineraries:
        for itinerary in past_itineraries:
            st.write(itinerary)
    else:
        st.write("No past itineraries yet.")

