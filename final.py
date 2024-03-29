import pickle
import re
from datetime import timedelta
import json
import math
import requests
import streamlit as st
from PIL import Image
from geopy.exc import GeocoderTimedOut
from geopy.geocoders import Nominatim
from openai import OpenAI
import logging
import os
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import nltk
from nltk import pos_tag, ne_chunk
from nltk.tokenize import word_tokenize
from amadeus import Client, ResponseError, Location
import zipfile
from dotenv import set_key, load_dotenv, dotenv_values
from docxtpl import DocxTemplate
from spire.doc import *
from spire.doc.common import *

RAPID_API_HOST = "booking-com.p.rapidapi.com"
st.session_state['data_changed'] = False
input_dict = {}
st.set_page_config(
    page_title="AI Tour Itinerary Generator",  # Set your desired title here
    page_icon="images/favicon.ico",  # Set path to your favicon image (.ico format)
)
st.title("Tour Itinerary Generator")

# Input fields for API keys
env_file_path = 'key.env'
env_vars = dotenv_values(env_file_path)

# Check if the key exists
if 'API_KEY' not in env_vars or 'RAPID_API_KEY' not in env_vars or 'AMADEUS_API_KEY' not in env_vars or 'AMADEUS_API_SECRET' not in env_vars or 'PEXELS_API_KEY' not in env_vars:
    st.subheader("Please enter API keys")
    API_KEY = st.text_input("Enter OpenAI API Key:")
    RAPID_API_KEY = st.text_input("Enter RapidAPI Key:")
    AMADEUS_API_KEY = st.text_input("Enter Amadeus client:")
    AMADEUS_API_SECRET = st.text_input("Enter Amadeus API secret:")
    PEXELS_API_KEY = st.text_input("Enter Pexels API key:")
    if st.button("Submit API Keys"):
        # Store the API keys in session state
        set_key(env_file_path, 'API_KEY', API_KEY)
        set_key(env_file_path, 'RAPID_API_KEY', RAPID_API_KEY)
        set_key(env_file_path, 'AMADEUS_API_KEY', AMADEUS_API_KEY)
        set_key(env_file_path, 'AMADEUS_API_SECRET', AMADEUS_API_SECRET)
        set_key(env_file_path, 'PEXELS_API_KEY', PEXELS_API_KEY)

        st.success("API Keys submitted successfully!")

else:
    load_dotenv('keys.env')
    API_KEY = os.getenv('API_KEY')
    RAPID_API_KEY = os.getenv("RAPID_API_KEY")
    AMADEUS_API_KEY = os.getenv("AMADEUS_API_KEY")
    AMADEUS_API_SECRET = os.getenv("AMADEUS_API_SECRET")
    PEXELS_API_KEY = os.getenv("PEXELS_API_KEY")

col1, col2 = st.columns(2)

input_dict['dest'] = col1.text_input("Destination", key='dest')
input_dict['src'] = col1.text_input("Source City", key='src')
input_dict['genre'] = col1.text_input("Genre", key='genre')
input_dict['type_of_travelers'] = col1.text_input("Type of Travelers", key='type', placeholder='ex. family, friends')
input_dict['mode_of_travel'] = col1.text_input("Mode of Travel", key='mode', placeholder='ex. flight, bus, train')
input_dict['num_days'] = col2.number_input("Number of Days", key='num_days', min_value=0, max_value=None, value=0,
                                           step=1, format="%d")
input_dict['start_date'] = col2.date_input("Start Date", key='start_date')
# Create sub-columns within col2
col21, col22 = col2.columns(2)

input_dict['num_adults'] = int(
    col21.number_input("Number of Adults", key='num_adults', min_value=0, max_value=None, value=0, step=1, format="%d"))
input_dict['num_children'] = int(
    col22.number_input("Number of Children", key='num_children', min_value=0, max_value=None, value=0, step=1,
                       format="%d"))
input_dict['price_per_person'] = col2.number_input("Price Per Person", key='price_per_person', min_value=0.0)
input_dict['average_age'] = col2.number_input("Average age", key='average_age', min_value=0, max_value=None, value=0,
                                              step=1, format="%d")
input_dict['food'] = 'non veg' if st.toggle('Include non-veg hotels') else 'veg'
special_note = st.text_area("Special Note(Optional)", key='special_note')


input_dict['num_tourists'] = input_dict['num_adults'] + input_dict['num_children']


client = OpenAI(api_key=API_KEY)

amadeus = Client(
    client_id='HlpBStOyyZ79qlc8cD4dTJEnsjnBv59Z',
    client_secret='VgElo0vAcc2QLiZ5'
)

function_descriptions = [
    {
        "name": "get_flight_hotel_info",
        "description": "Find the flights between cities and hotels within cities for residence.",
        "parameters": {
            "type": "object",
            "properties": {
                "loc_list": {
                    "type": "array",
                    "items": {
                        "type": "string"
                    },
                    "description": "The ordered list of names of cities in the tour. e.g. ['Mumbai', 'Paris']"
                },

                "date_list": {
                    "type": "array",
                    "items": {
                        "type": "string"
                    },
                    "description": "The ordered list of dates for arrival in the cities in YYYY-MM-DD format."
                },

                "iata_list":{
                    "type": "array",
                    "items": {
                        "type": "string"
                    },
                    "description": "The ordered list of IATA codes of cities in the tour which you have selected. e.g. ['BOM' for 'Mumbai', 'PAR' for 'Paris']"
                }

            },
            "required": ["loc_list", "date_list", "iata_list"],
        },
    }
]

# Set the logging level (DEBUG for most details)
logging.basicConfig(level=logging.DEBUG)

# Create a logger for your specific API usage
logger = logging.getLogger('booking_com_api')


def make_booking_com_api_call(url, headers, data):
    logger.debug("API Request:")
    logger.debug(f"URL: {url}")
    logger.debug(f"Headers: {headers}")
    logger.debug(f"Body: {data}")

    # Make your API call using requests library
    response = requests.get(url, headers=headers, json=data)

    logger.debug("API Response:")
    logger.debug(f"Status Code: {response.status_code}")
    logger.debug(f"Headers: {response.headers}")
    logger.debug(f"Response Body: {response.json()}")

    return response


def get_hotel_data(city, checkin_date, checkout_date, num_adults, num_children):
    city_dict = {}

    # Geocode city to get latitude and longitude
    def geocode_with_retry(location):
        geolocator = Nominatim(user_agent="my_geocoder")
        max_retries = 3

        for attempt in range(max_retries):
            try:
                geocoded_location = geolocator.geocode(location)
                return geocoded_location.raw
            except GeocoderTimedOut as e:
                print(f'Attempt {attempt + 1} failed: {e}')
                if attempt < max_retries - 1:
                    print(f'Retrying...')
                    continue

        print('Max retries exceeded. Unable to geocode the location.')
        return None
    location = geocode_with_retry(city)
    print('Location', location)

    if location:
        lat = location['lat']
        long = location['lon']

        # Define URL and query parameters for Booking.com API
        url = "https://booking-com.p.rapidapi.com/v1/hotels/search-by-coordinates"
        num_rooms = math.ceil((num_adults + num_children) / 3)
        querystring = {
            "locale": "en-gb",
            "room_number": num_rooms,
            "checkin_date": checkin_date,
            "checkout_date": checkout_date,
            "filter_by_currency": "INR",
            "longitude": long,
            "latitude": lat,
            "adults_number": num_adults,
            "order_by": "popularity",
            "units": "metric",
            "page_number": "0",
        }
        if num_children > 0:
            querystring["children_number"] = num_children
        headers = {
            "X-RapidAPI-Key": RAPID_API_KEY,
            "X-RapidAPI-Host": RAPID_API_HOST
        }

        # Send request to Booking.com API
        response = requests.get(url, headers=headers, params=querystring)
        if response.status_code == 200:
            data = response.json()
            if "result" in data:
                city_dict[city] = []
                st.write(f"Finding hotels for {city}")
                # st.write("Hotel search results for", city, ":")
                for index, hotel in enumerate(data["result"], start=1):
                    hotel_dict = {}
                    hotel_dict['hotel_name'] = hotel.get("hotel_name", "N/A")
                    price = hotel.get("min_total_price", "N/A")
                    if price is not None:
                        hotel_dict['price'] = str(price)
                    else:
                        hotel_dict['price'] = "N/A"
                    hotel_dict['address'] = hotel.get("address", "N/A")
                    hotel_dict['rating'] = hotel.get("review_score", "N/A")
                    # st.write(f"{index}. {hotel_dict['hotel_name']} - Address: {hotel_dict['address']}, Price for one day: {hotel_dict['price']} INR, Rating: {hotel_dict['rating']}")
                    city_dict[city].append(hotel_dict)
            else:
                st.write(f"No hotel results found for {city}.")
        else:
            st.write(f"Failed to retrieve hotel search results for {city}.")
    else:
        st.write(f"Invalid city name: {city}")

    return city_dict


def flight_search(input_dict):
    num_adults = input_dict['num_adults']
    cities = input_dict['cities']
    dates = input_dict['dates']
    flights = {}

    # Convert city names to IATA codes
    iata_codes = input_dict['iata_codes']
    # for city in cities:
    #     try:
    #         response = amadeus.reference_data.locations.get(
    #             keyword=city,
    #             subType=Location.ANY
    #         )
    #         if response.data:  # Check if response data is not empty
    #             iata_code = response.data[0]['iataCode']
    #             iata_codes[city] = iata_code
    #             print(city, iata_code)
    #         else:
    #             print(f"Error: Could not find airport code for {city}. No matching locations found.")
    #     except ResponseError as error:
    #         print(f"Error: An error occurred while searching for airport code for {city}: {error}")

    for i in range(len(cities) - 1):
        src = cities[i]
        dest = cities[i + 1]
        flights[src] = []
        src_iata = iata_codes[i]
        dest_iata = iata_codes[i + 1]

        try:
            # Check if the source city exists in the iata_codes dictionary
            # if src in iata_codes:
            #     src_iata = iata_codes[i]
            # else:
            #     print(f"Error: Could not find airport code for source city {src}.")
            #     continue
            #
            # # Check if the destination city exists in the iata_codes dictionary
            # if dest in iata_codes:
            #     dest_iata = iata_codes[i+1]
            # else:
            #     print(f"Error: Could not find airport code for destination city {dest}.")
            #     continue

            response = amadeus.shopping.flight_offers_search.get(
                originLocationCode=src_iata,
                destinationLocationCode=dest_iata,
                departureDate=dates[i],
                adults=min(num_adults,9),
                travelClass='ECONOMY',
                currencyCode='INR',
                max=20
            )

            flights[src].append(response.result['data'])

        except ResponseError as error:
            print(f"Error: List of flights for {src} to {dest} is not available.")
            print(error.description())

    return flights


def display_flight_info(flight_data):
    city_flight_info = {}  # Dictionary to store flight info by city

    for src, flights in flight_data.items():
        city_flight_info[src] = []  # Initialize list for each city
        for flight_list in flights:
            for flight in flight_list:
                # Extract specific information from the flight data
                departure_time = flight['itineraries'][0]['segments'][0]['departure']['at']
                arrival_time = flight['itineraries'][0]['segments'][-1]['arrival']['at']
                airline = flight['validatingAirlineCodes'][0]
                price = flight['price']['grandTotal']
                currency = flight['price']['currency']

                # Format flight information as a dictionary
                flight_info = {
                    "Airline": airline,
                    "Departure Time": departure_time,
                    "Arrival Time": arrival_time,
                    "Price": price,
                    "Currency": currency
                }
                # Append flight information to the list for the corresponding city
                city_flight_info[src].append(flight_info)

    return city_flight_info


@st.cache_data(show_spinner=False)
def generate_itinerary(input_dict):
    # Part 1: generate the list of cities and get the hotels
    # Call the OpenAI API for creating the list of cities and dates
    input_dict['end_date'] = str(input_dict['start_date'] + timedelta(days=input_dict['num_days']))
    user_prompt = f"Generate a list of cities for a tour of {input_dict['dest']} for {input_dict['num_tourists']} " \
                  f"people with {input_dict['num_adults']} adults, purpose as {input_dict['genre']} " \
                  f"for {input_dict['num_days']} days and a budget per person of {input_dict['price_per_person']} INR starting on " \
                  f"{input_dict['start_date']}, ending on {input_dict['end_date']}. Call the function 'get_flight_hotel_info'"

    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": user_prompt}],
        # Add function calling
        functions=function_descriptions,
        function_call="auto",  # specify the function call
        max_tokens=200
    )
    output = completion.choices[0].message
    cities = json.loads(output.function_call.arguments).get("loc_list")
    dates = json.loads(output.function_call.arguments).get("date_list")
    iata_codes = json.loads(output.function_call.arguments).get("iata_list")
    dates.append(input_dict['end_date'])

    print(cities)
    print(dates)
    print(iata_codes)

    input_dict['cities'] = cities
    input_dict['dates'] = dates
    input_dict['iata_codes'] = iata_codes
    all_city_dict = {}

    printables = {}
    city_string = ''
    for city in cities:
        city_string += city + ' - '
    st.subheader("Cities: ")
    st.write(city_string)
    printables['city_string'] = city_string

    for i in range(len(cities)):
        # st.write(cities[i], dates[i], dates[i+1], input_dict['num_adults'], input_dict['num_children'])
        all_city_dict.update(
            get_hotel_data(cities[i], dates[i], dates[i + 1], input_dict['num_adults'], input_dict['num_children']))
    input_dict['hotels_by_city'] = all_city_dict

    # Part 2: Actually generate the itinerary
    user_message = f"Design a detailed itinerary for a trip from {input_dict['src']} to {input_dict['dest']} starting from {input_dict['start_date']} and for " \
                   f"{input_dict['num_days']} days. The ordered list of cities is {cities} and of dates is {dates}. The budget for this trip is {input_dict['price_per_person']} INR per person. This trip is designed " \
                   f"for {input_dict['num_tourists']} mainly with their {input_dict['type_of_travelers']} with an average age of {input_dict['average_age']}.The " \
                   f"primary interests for activities are {input_dict['genre']}.The preferred mode(s) of travel include " \
                   f"{input_dict['mode_of_travel']}.The group prefers {input_dict['food']} food. Please structure the itinerary with a detailed " \
                   f"plan for each day along with a every day title and no word should get repeated in a title of other days, including activities, locations, weather according to the season they are " \
                   f"travelling and estimated travel distances and times(Do not give null values if you cannot extract information). Write the travel time and distance in the day's subheading. " \
                   f"Ensure to consider the preferences and " \
                   f"interests of the group for each day's schedule. Important considerations: Factor in travel time " \
                   f"between destinations. Suggest local transportation options. Include a mix of activities that cater" \
                   f" to the group's interests. Also add distance of travel for each day and approx time(do not give null value sif not available) " \
                   f"of travel. Also you can give a name for each day in the itinerary which will be more " \
                   f"appealing. Keep the response descriptive and . Give a title to the itinerary but make sure you don't repeat location names in multiple days also you can mention prime locations in title that are going to be there in iternary. Do not suggest any activities " \
                   f"in the first city if the travel time and distance is more otherwise we can suggest activities." \
                   f"Finally the description for each day which should look like if a human is speaking(this paragraph will be under the heading for each day)" \
 \
    # Generate the travel itinerary using the modified user message
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": user_message,
            }
        ],
        model="gpt-3.5-turbo",
        stream=True,
    )

    st.subheader("Itinerary")
    response = st.write_stream(chat_completion)

    flight_data = flight_search(input_dict)

    # Display flight information
    flight_info = display_flight_info(flight_data)
    content = response

    # Split content into individual days
    days = content.split("\n\n")
    titles_and_days = re.findall(r"Day (\d+): (.+?)(?=\n)", content)

    print(titles_and_days)  # Debugging: print titles_and_days to see its structure

    # Check if the directory exists, if not, create it
    if not os.path.exists('images'):
        os.makedirs('images')

    # Iterate over titles and fetch images
    for day_number, title in titles_and_days:
        image = fetch_image(title)
        if image:
            image.save(f"images/day{day_number}.jpg")
            print(f"Image saved successfully for {title.strip()} as day{day_number}.jpg")
        else:
            print(f"No image found for {title.strip()}")

    st.session_state['input_dict'] = input_dict
    with open('input_dict.pickle', 'wb') as handle:
        pickle.dump(input_dict, handle, protocol=pickle.HIGHEST_PROTOCOL)
    return response, all_city_dict, flight_info, days, city_string


def fetch_image(title):
    # Define your Pexels API key
    API_KEY = PEXELS_API_KEY

    # Define the search query using the title
    query = title

    # Define the Pexels API endpoint
    endpoint = f"https://api.pexels.com/v1/search?query={query}&per_page=1"  # Fetch only one image per title

    # Define headers including your API key
    headers = {
        'Authorization': API_KEY
    }

    try:
        # Send a GET request to the Pexels API
        response = requests.get(endpoint, headers=headers)

        # Check if the request was successful (status code 200)
        if response.status_code == 200:
            # Parse the JSON response
            data = response.json()

            # Extract the list of photos from the response
            photos = data['photos']

            if photos:
                # Extract image URL from the first photo
                image_url = photos[0]['src']['original']

                # Download the image
                image_response = requests.get(image_url)

                # Open the image using PIL
                image = Image.open(BytesIO(image_response.content))

                return image
            else:
                print("No images found for", title)
                return None

        else:
            print("Failed to fetch images. Status code:", response.status_code)
            return None

    except Exception as e:
        print("An error occurred:", str(e))
        return None


def fetch_images_when_selected(selected_proper_noun):
    if selected_proper_noun:
        fetch_images_from_pexels([selected_proper_noun])


def fetch_images_from_pexels(proper_noun):
    # Pexels API configuration
    API_KEY = PEXELS_API_KEY
    base_url = 'https://api.pexels.com/v1/search'

    # Set up parameters for the request
    params = {'query': proper_noun, 'per_page': 1}

    # Make the request to Pexels API
    response = requests.get(base_url, params=params, headers={'Authorization': API_KEY})
    print("Fetching images for:", proper_noun)

    if response.status_code == 200:
        # Extract image URL from the response
        data = response.json()
        if data['total_results'] > 0:
            image_url = data['photos'][0]['src']['large']
            # Save the image to a file
            save_image(image_url, proper_noun)
        else:
            print(f"No images found for {proper_noun}")
    else:
        print(f"Error fetching images for {proper_noun}")


def save_image(image_url, proper_nouns):
    # Create images directory if not exists
    if not os.path.exists('images'):
        os.makedirs('images')

    # Download and save the image
    response = requests.get(image_url)
    if response.status_code == 200:
        with open(f'images/{proper_nouns}.jpg', 'wb') as f:
            f.write(response.content)
            print(f"Image saved for {proper_nouns}")
    else:
        print(f"Failed to download image for {proper_nouns}")


# Download required NLTK resources
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('maxent_ne_chunker')
nltk.download('words')


def extract_proper_nouns(text):
    # Tokenize the text into words
    words = word_tokenize(text)

    # Perform part-of-speech tagging
    tagged_words = pos_tag(words)

    # Perform named entity recognition (NER)
    ne_tree = ne_chunk(tagged_words)

    # Extract proper nouns from NER results
    proper_nouns = []
    for subtree in ne_tree:
        if isinstance(subtree, nltk.Tree) and subtree.label() == 'PERSON':
            # If it's a PERSON entity, treat it as a proper noun
            proper_nouns.append(" ".join([word for word, _ in subtree.leaves()]))
        elif isinstance(subtree, tuple) and subtree[1] == 'NNP':
            # If it's tagged as a proper noun (NNP), add it to the list
            proper_nouns.append(subtree[0])

    return proper_nouns


@st.cache_data(show_spinner=False)
def text_to_doc(itinerary, input_dict):
    # document = Document()
    # paragraph = document.add_paragraph()
    # run = paragraph.add_run()
    # run.add_picture("icons/logo.png", width=Inches(2.7))  # Adjust width as needed
    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    # first_line = itinerary.split('\n')[0]
    #
    # # Add the first line as a centered header
    # header = document.add_heading(level=0)
    # header_run = header.add_run(first_line)
    # header_run.font.size = Pt(22)
    # header_run.font.name = 'Bahnschrift'
    # header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    # # Add subheader with small images
    # subheader_text = f"{input_dict['num_days']} | {input_dict['start_date']} | {input_dict['dest']} to {input_dict['src']}"
    # subheader_paragraph = document.add_paragraph()
    # subheader_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    # # Function to add image and text with specified width
    # def add_image_run(paragraph, image_path, text, image_width):
    #     run = paragraph.add_run()
    #     run.add_picture(image_path, width=image_width)
    #     run.add_text(text)
    #
    # # Define a list of different image paths
    # image_paths = ["icons/cal.png", "icons/global.jpeg", "icons/gps.png"]  # Add more paths as needed
    #
    # # Add subheader components with images
    # for index, component in enumerate(subheader_text.split(" | ")):
    #     image_path = image_paths[index] if index < len(image_paths) else image_paths[
    #         -1]  # Use the last image if there are more components than images
    #     add_image_run(subheader_paragraph, image_path, component, Inches(0.15))
    #
    # paragraph3 = document.add_paragraph()
    # run3 = paragraph3.add_run("Main Focus ")
    # run3.bold = True
    # run3 = paragraph3.add_run(input_dict['genre'])
    #
    # paragraph4 = document.add_paragraph()
    # run4 = paragraph4.add_run("Commences on: ")
    # run4.bold = True
    # run4 = paragraph4.add_run(str(input_dict['start_date']))
    #
    # paragraph6 = document.add_paragraph()
    # run6 = paragraph6.add_run("Budget : ")
    # run6.bold = True
    # run6 = paragraph6.add_run(str(input_dict['price_per_person']))
    #
    # # Add countries and cities covered
    # paragraph5 = document.add_paragraph()
    # run5 = paragraph5.add_run("Countries and Cities Covered:")
    # run5.bold = True
    #
    # # Join the city names with a comma
    # city_names = ", ".join(input_dict['cities'])
    #
    # # Add all city names in a single paragraph
    # document.add_paragraph(f"- {city_names}")
    #
    # line_paragraph = document.add_paragraph()
    # line_run = line_paragraph.add_run(
    #     "_")
    #
    # # Set the font size of the line
    # line_font = line_run.font
    # line_font.size = Pt(12)
    #
    # # Set the paragraph alignment to center
    # line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    # # Function to check if a specific word is present in the itinerary
    # paragraph = document.add_paragraph("")
    #
    # paragraph.bold = True
    #
    # def contains_word(itinerary, word):
    #     return word.lower() in itinerary.lower()
    #
    # # Function to add image without text
    # def add_image_without_text(paragraph, image_path, image_width):
    #     run = paragraph.add_run()
    #     run.add_picture(image_path, width=image_width)
    #
    # # Check if the word "hotel" is present in the itinerary
    # if contains_word(itinerary, "hotel"):
    #     add_image_without_text(paragraph, "icons/hotel.png", Inches(1))
    # if contains_word(itinerary, "flight"):
    #     add_image_without_text(paragraph, "icons/flight.jpeg", Inches(1))
    # if contains_word(itinerary, "eat"):
    #     add_image_without_text(paragraph, "icons/meal.png", Inches(1))
    # if contains_word(itinerary, "attraction"):
    #     add_image_without_text(paragraph, "icons/sightseeing.png", Inches(1))
    # if contains_word(itinerary, "train"):
    #     add_image_without_text(paragraph, "icons/train.jpeg", Inches(1))
    # if contains_word(itinerary, "religious"):
    #     add_image_without_text(paragraph, "icons/religious.png", Inches(1))
    # if contains_word(itinerary, "work"):
    #     add_image_without_text(paragraph, "icons/work.png", Inches(1))
    #
    # paragraph.add_run().add_break()  # Add a blank line
    #
    # paragraph.add_run().add_break()  # Add a blank line
    #
    # paragraph5 = document.add_paragraph()
    # a = paragraph5.add_run("Tour Itinerary :")
    # a.bold = True
    # for line in itinerary.split('\n')[2:]:
    #     paragraph = document.add_paragraph()
    #     image_added = False  # Flag variable to track whether an image has been added to the line
    #     for char in line:
    #         # if "Dinner" in line and "Day" not in line and not image_added:
    #         #     add_image_without_text(paragraph, "icons/meal.png", Inches(0.4))
    #         #     image_added = True
    #         # elif "Lunch" in line and "Day" not in line and not image_added:
    #         #     add_image_without_text(paragraph, "icons/meal.png", Inches(0.4))
    #         #     image_added = True
    #         # elif "Breakfast" in line and "Day" not in line and not image_added:
    #         #     add_image_without_text(paragraph, "icons/breakfast.jpeg", Inches(0.4))
    #         #     image_added = True
    #         if "Travel Distance" in line and "Day" not in line and not image_added:
    #             add_image_without_text(paragraph, "icons/travel_distance.jpeg", Inches(0.4))
    #             image_added = True
    #         elif "Travel Time" in line and "Day" not in line and not image_added:
    #             add_image_without_text(paragraph, "icons/travel_time.jpeg", Inches(0.4))
    #             image_added = True
    #         # elif "Fly" in line and "Day" not in line and not image_added:
    #         #     add_image_without_text(paragraph, "icons/flight.jpeg", Inches(0.4))
    #         #     image_added = True
    #
    #     for char in line:
    #         if "Day" in line:
    #             run = paragraph.add_run(char)
    #             run.bold = True
    #             run.font.size = Pt(14)
    #             run.font.name = 'Aptos'
    #
    #
    #         else:
    #             run.font.name = 'Avenir Next LT Pro'
    #             run.font.size = Pt(12)
    #             run = paragraph.add_run(char)
    #
    #     # Add image after each day's description
    #     if line.startswith("Day"):
    #         day_number = line.split()[1].rstrip(":")  # Extract day number and remove the colon
    #         image_path = f"images/day{day_number}.jpg"  # Form the image path
    #         if os.path.exists(image_path):
    #             paragraph.add_run("\n")
    #             add_image_without_text(paragraph, image_path, Inches(2))  # Adjust width as needed
    #
    # # Add a table
    # table_data = [
    #     ["Destination", " Hotel"],
    # ]
    # for city in input_dict['cities']:
    #     table_data.append([city, ''])
    #
    # table = document.add_table(rows=len(table_data), cols=2)
    #
    # # adding data to table
    # for i, row_data in enumerate(table_data):
    #     for j, cell_data in enumerate(row_data):
    #         table.cell(i, j).text = cell_data
    #
    # # Apply alignment to the table
    # for row in table.rows:
    #     for cell in row.cells:
    #         for paragraph in cell.paragraphs:
    #             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    # # Define a custom table style (optional)
    # table.style = 'Table Grid'
    day_itineraries = generate_day_itineraries(itinerary)
    city_names = ", ".join(input_dict['cities'])
    folder_name = "generated_itineraries"

    # Delete the previously generated documents
    for filename in os.listdir(folder_name):
        file_path = os.path.join(folder_name, filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Error deleting {file_path}: {e}")

    first_page = DocxTemplate('mergeDocs/front_page.docx')
    context = {
        'tour_heading': itinerary.split('\n')[0],
        'num_days': input_dict['num_days'],
        'budget': input_dict['price_per_person'],
        'cities': city_names,
    }
    first_page.render(context)
    file_path = os.path.join(folder_name, 'cover_page.docx')
    first_page.save(file_path)

    # Load the template document
    tpl = DocxTemplate("mergeDocs/daywise_itinerary.docx")

    for day_number, day_itinerary in day_itineraries.items():
        # Extract the first line of the itinerary
        first_line = day_itinerary.split('\n')[0]
        print("Inside the text_to_doc func: ", day_itinerary, 'First line: ', first_line, 'Second line: ', day_itinerary)
        # Join city names into a comma-separated string
        first_newline_index = day_itinerary.find('\n')

        # Check if '\n' exists in the string
        if first_newline_index != -1:
            # Extract the substring starting from the index after the first '\n'
            day_itinerary = day_itinerary[first_newline_index + 1:]

        # Define the context dictionary
        context = {
            'tour_heading': first_line,
            'num_days': input_dict['num_days'],
            'budget': input_dict['price_per_person'],
            'day_itinerary': day_itinerary,
            'day_title': first_line
        }

        # Replace placeholders in the document
        tpl.render(context)

        # Replace the placeholder for day in the itinerary and change day_title
        # for paragraph in tpl.paragraphs:
        #     if '{{day_title}}' in paragraph.text:
        #         paragraph.text = paragraph.text.replace('{{day_title}}', f'Day {day_number}')
        #     if '{{num_days}}' in paragraph.text:
        #         paragraph.text = paragraph.text.replace('{{num_days}}', str(input_dict['num_days']))
        #     if '{{cities}}' in paragraph.text:
        #         paragraph.text = paragraph.text.replace('{{cities}}', city_names)
        #     if '{{budget}}' in paragraph.text:
        #         paragraph.text = paragraph.text.replace('{{budget}}', input_dict['price_per_person'])

        # Create a folder to store the generated documents

        os.makedirs(folder_name, exist_ok=True)

        # Modify the file path where the documents are saved
        file_path = os.path.join(folder_name, f'day_{day_number}_itinerary.docx')
        tpl.save(file_path)

    # Create a Document object
    destDoc = Document()
    # # Load the destination document
    # destDoc.LoadFromFile("mergeDocs/front_page.docx")

    # Define the folder path containing the files to merge
    folder_path = "generated_itineraries"

    # Create the directory if it doesn't exist
    os.makedirs(folder_path, exist_ok=True)

    # List all files in the folder
    files_to_merge = os.listdir(folder_path)

    # Filter only the .docx files
    files_to_merge = [file for file in files_to_merge if file.endswith('.docx')]

    # Loop through the list
    for file in files_to_merge:
        # Construct the full file path
        file_path = os.path.join(folder_path, file)

        # Load the source document
        sourceDoc = Document()
        sourceDoc.LoadFromFile(file_path)

        # Keep the formatting of the source document when it is merged
        # sourceDoc.KeepSameFormat = True

        # Import the content from the document into the destination document
        destDoc.ImportContent(sourceDoc)

    # Save the result document
    destDoc.SaveToFile("Itinerary.docx", FileFormat.Docx2016)
    destDoc.Close()
    sourceDoc.Close()

    with open("Itinerary.docx", "rb") as file:
        bytes_content = file.read()
    return bytes_content


@st.cache_data(show_spinner=False)
def display_image_choices(days):
    proper_nouns_by_day = {}
    for day in days:
        lines = day.split("\n")
        day_name = lines[0]
        day_content = "\n".join(lines[1:])  # Exclude the day name

        # Extract keywords for the day
        keywords = extract_proper_nouns(day_content)
        unique_proper_nouns = list(set(keywords))

        # Store keywords for the day
        proper_nouns_by_day[day_name] = unique_proper_nouns
    st.session_state['proper_nouns_by_day'] = proper_nouns_by_day
    return proper_nouns_by_day

@st.cache_data(show_spinner=False)
def create_zip_file():
    # Get the list of image files in the images directory
    image_files = [filename for filename in os.listdir("images") if filename.endswith(".jpg")]

    # Create a zip file
    with zipfile.ZipFile("images.zip", "w") as zipf:
        for image_file in image_files:
            zipf.write(os.path.join("images", image_file), image_file)

    with open("images.zip", "rb") as f:
        zip_data = f.read()

    return zip_data

def get_day_itinerary(itinerary, day_number):
    # Split the itinerary into days
    days = itinerary.split("Day ")
    for day in days[1:]:
        if day.startswith(str(day_number) + ":"):
            day = day.replace('*', '')
            day = day.replace('###', '')
            return "Day " + day

def generate_day_itineraries(itinerary):
    day_itineraries = {}
    # Get the maximum day number mentioned in the itinerary
    max_day = max([int(day.split(":")[0]) for day in itinerary.split("Day ")[1:]])
    # Generate day itineraries for each day
    for day_number in range(1, max_day + 1):
        day_itinerary = get_day_itinerary(itinerary, day_number)
        if day_itinerary:
            print('Inside generate_day_itinerary: \n', day_itinerary)
            day_itineraries[day_number] = day_itinerary
    return day_itineraries



################## Main Code #################
if st.session_state.get('input_dict', False):
    for key in input_dict.keys():
        if input_dict[key] != st.session_state['input_dict'][key]:
            st.session_state['data_changed'] = True
            break

if st.button("Generate Itinerary", type="primary"):
    null_flag = False
    if 'OPENAI_API_KEY' not in env_vars or 'X-RapidAPI-Key' not in env_vars or 'AMADEUS_API_KEY' not in env_vars or 'AMADEUS_API_SECRET' not in env_vars or 'PEXELS_API_KEY' not in env_vars:
        st.warning('Enter all the API keys')
        null_flag = True
    for key in input_dict.keys():
        if input_dict[key] is None:
            st.warning(f'Please enter {key}!')
            null_flag = True
            break

    if not null_flag:
        generated_itinerary, city_dict, flight_info, days, city_string = generate_itinerary(input_dict)
        st.session_state["cached_data_generated"] = True
        st.session_state['data_changed'] = False
        isGenerated = True

elif st.session_state.get("cached_data_generated", False) and not st.session_state['data_changed']:
    generated_itinerary, city_dict, flight_info, days, city_string = generate_itinerary(input_dict)

if st.session_state.get("cached_data_generated", False) and not st.session_state['data_changed']:
    st.subheader("Hotels")
    for city, hotels in city_dict.items():
        city_expander = st.expander(f"{city}")
        with city_expander:
            for hotel in hotels:
                st.write(f"- {hotel['hotel_name']}")
                st.write(f"  Address: {hotel['address']}")
                st.write(f"  Price per day: {hotel['price']} INR")
                st.write(f"  Rating: {hotel['rating']}")
                # Add more details as needed (amenities, images, etc.)
                st.write("---")  # Separator between hotels\

    st.subheader("Flight Details")
    for city, flights in flight_info.items():
        city_expander = st.expander(f"{city}")
        with city_expander:
            for flight in flights:
                st.write(f"- {flight['Airline']}")
                st.write(f"  Departure Time: {flight['Departure Time']}")
                st.write(f"  Arrival Time: {flight['Arrival Time']}")
                st.write(f"  Price: {flight['Price']} INR")
                # Add more details as needed (amenities, images, etc.)
                st.write("---")


    doc_io = text_to_doc(generated_itinerary, st.session_state['input_dict'])

    st.download_button(
        label="Download Word Document",
        data=doc_io,
        file_name=f"{input_dict['dest']} Itinerary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# Change icons
# Use a single image for the poster
# Mention the meaning of each icon below the title
# Mention end date
# Number of nights per city
# Date of each day
# change the font
